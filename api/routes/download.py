"""
api/routes/download.py

GET /download-notes — download highlights as PDF, TXT, or DOCX.
"""

import os
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

from auth.dependencies import get_current_user
from auth.db import get_meeting_name

router = APIRouter()


@router.get("/download-notes")
def download_notes(
    meeting_id: str,
    format: str = "pdf",
    user: dict = Depends(get_current_user)
):

    # ── resolve meeting name and verify ownership ─────────
    meeting_name = get_meeting_name(meeting_id, user["user_id"])
    if not meeting_name:
        raise HTTPException(status_code=403, detail="Meeting not found or access denied")

    meeting_name = "".join(
        c for c in meeting_name if c.isalnum() or c in (" ", "-", "_")
    ).strip()

    download_time = datetime.now().strftime("%d %b %Y, %I:%M %p")
    txt_path = f"Notes/highlights_{meeting_id}.txt"

    if not os.path.exists(txt_path):
        return {"error": "Highlights not generated yet."}

    with open(txt_path, "r", encoding="utf-8") as f:
        text = f.read()

    # ── PDF ───────────────────────────────────────────────
    if format == "pdf":
        pdf_path = f"Notes/{meeting_name}_{meeting_id}.pdf"
        doc = SimpleDocTemplate(pdf_path)
        styles = getSampleStyleSheet()
        elements = [
            Paragraph("Meeting Highlights", styles["Heading1"]),
            Spacer(1, 20),
            Paragraph(f"Meeting: {meeting_name}", styles["Normal"]),
            Paragraph(f"Downloaded: {download_time}", styles["Normal"]),
            Spacer(1, 20),
        ]
        for line in text.split("\n"):
            elements.append(Paragraph(line, styles["BodyText"]))
            elements.append(Spacer(1, 8))
        doc.build(elements)
        return FileResponse(pdf_path, media_type="application/pdf",
                            filename=f"{meeting_name}.pdf")

    # ── TXT ───────────────────────────────────────────────
    elif format == "txt":
        header = f"Meeting: {meeting_name}\nDownloaded: {download_time}\n\n"
        temp_path = f"Notes/{meeting_name}_{meeting_id}.txt"
        with open(temp_path, "w", encoding="utf-8") as f:
            f.write(header + text)
        return FileResponse(temp_path, media_type="text/plain",
                            filename=f"{meeting_name}.txt")

    # ── DOCX ──────────────────────────────────────────────
    elif format == "docx":
        docx_path = f"Notes/{meeting_name}_{meeting_id}.docx"
        document = Document()
        document.add_heading("Meeting Highlights", 0)
        document.add_paragraph(f"Meeting: {meeting_name}")
        document.add_paragraph(f"Downloaded: {download_time}")
        document.add_paragraph("")
        for line in text.split("\n"):
            document.add_paragraph(line)
        document.save(docx_path)
        return FileResponse(
            docx_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{meeting_name}.docx"
        )

    else:
        return {"error": "Invalid format"}
